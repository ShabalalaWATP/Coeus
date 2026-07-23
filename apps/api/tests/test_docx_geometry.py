from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document

from coeus.services import document_extraction, docx_geometry
from coeus.services.document_extraction import DocumentExtractionError, extract_pages

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
WORD = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def test_docx_geometry_accepts_exact_row_and_document_limits() -> None:
    document_extraction._validate_docx_table_geometry(
        _geometry_xml([[docx_geometry.MAX_DOCX_LOGICAL_CELLS_PER_ROW]])
    )
    document_extraction._validate_docx_table_geometry(_geometry_xml([[1_000]] * 10))
    explicit_zero_offset = _geometry_xml([[1]]).replace(
        b"<w:tr>",
        b'<w:tr><w:trPr><w:gridBefore w:val="0"/></w:trPr>',
    )
    document_extraction._validate_docx_table_geometry(explicit_zero_offset)


def test_docx_geometry_rejects_row_and_document_limit_overflow() -> None:
    with pytest.raises(DocumentExtractionError, match="docx_table_geometry_too_large"):
        document_extraction._validate_docx_table_geometry(
            _geometry_xml([[docx_geometry.MAX_DOCX_LOGICAL_CELLS_PER_ROW + 1]])
        )

    with pytest.raises(DocumentExtractionError, match="docx_table_geometry_too_large"):
        document_extraction._validate_docx_table_geometry(_geometry_xml([[1_000]] * 10 + [[1]]))


def test_docx_geometry_rejects_unsafe_xml_and_invalid_merge_value() -> None:
    with pytest.raises(DocumentExtractionError, match="docx_parse_failed"):
        document_extraction._validate_docx_table_geometry(b"<w:document")

    invalid_merge = _vertical_geometry_xml(1).replace(b'w:val="restart"', b'w:val="sideways"')
    with pytest.raises(DocumentExtractionError, match="docx_table_geometry_invalid"):
        document_extraction._validate_docx_table_geometry(invalid_merge)


@pytest.mark.parametrize("span", ["0", "-1", "not-a-number", None])
def test_docx_geometry_rejects_invalid_grid_spans(span: str | None) -> None:
    with pytest.raises(DocumentExtractionError, match="docx_table_geometry_invalid"):
        document_extraction._validate_docx_table_geometry(_geometry_xml([[span]]))


def test_docx_rejects_grid_span_before_python_docx_materialises_cells(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    content = _docx_with_grid_span(docx_geometry.MAX_DOCX_LOGICAL_CELLS_PER_ROW + 1)

    def unexpected_document(*_args: object, **_kwargs: object) -> object:
        raise AssertionError("python-docx must not receive over-budget table geometry")

    monkeypatch.setattr("docx.Document", unexpected_document)

    with pytest.raises(DocumentExtractionError, match="docx_table_geometry_too_large"):
        extract_pages(content, DOCX_MIME)


def test_docx_geometry_bounds_vertical_merge_depth_and_materialisation_work(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document_extraction._validate_docx_table_geometry(
        _vertical_geometry_xml(docx_geometry.MAX_DOCX_VERTICAL_MERGE_DEPTH)
    )
    with pytest.raises(DocumentExtractionError, match="docx_table_geometry_too_large"):
        document_extraction._validate_docx_table_geometry(
            _vertical_geometry_xml(docx_geometry.MAX_DOCX_VERTICAL_MERGE_DEPTH + 1)
        )

    monkeypatch.setattr(docx_geometry, "MAX_DOCX_CELL_MATERIALISATION_WORK", 6)
    document_extraction._validate_docx_table_geometry(_vertical_geometry_xml(3))
    with pytest.raises(DocumentExtractionError, match="docx_table_geometry_too_large"):
        document_extraction._validate_docx_table_geometry(_vertical_geometry_xml(4))


def test_docx_rejects_deep_vertical_merge_before_python_docx_materialises_cells(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    content = _docx_with_vertical_merge(docx_geometry.MAX_DOCX_VERTICAL_MERGE_DEPTH + 1)

    def unexpected_document(*_args: object, **_kwargs: object) -> object:
        raise AssertionError("python-docx must not receive over-budget vertical merges")

    monkeypatch.setattr("docx.Document", unexpected_document)
    with pytest.raises(DocumentExtractionError, match="docx_table_geometry_too_large"):
        extract_pages(content, DOCX_MIME)


def test_docx_extraction_preserves_legitimate_vertical_merge() -> None:
    content = _docx_with_vertical_merge(10)

    pages = extract_pages(content, DOCX_MIME)

    assert "Synthetic vertical assessment" in pages[0].text


def test_docx_extraction_preserves_legitimate_merged_cells() -> None:
    stream = BytesIO()
    document = Document()
    table = document.add_table(rows=1, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Synthetic merged assessment"
    document.save(stream)

    pages = extract_pages(stream.getvalue(), DOCX_MIME)

    assert len(pages) == 1
    assert "Synthetic merged assessment" in pages[0].text


def test_docx_rejects_ambiguous_duplicate_document_parts() -> None:
    source = BytesIO()
    document = Document()
    document.add_paragraph("Synthetic")
    document.save(source)
    with ZipFile(BytesIO(source.getvalue())) as archive:
        members = [(member.filename, archive.read(member)) for member in archive.infolist()]
    output = BytesIO()
    with (
        pytest.warns(UserWarning, match="Duplicate name"),
        ZipFile(output, "w", ZIP_DEFLATED) as archive,
    ):
        for name, value in members:
            archive.writestr(name, value)
        archive.writestr("word/document.xml", members[0][1])

    with pytest.raises(DocumentExtractionError, match="docx_not_extractable"):
        extract_pages(output.getvalue(), DOCX_MIME)


def _geometry_xml(rows: list[list[int | str | None]]) -> bytes:
    encoded_rows: list[str] = []
    for row in rows:
        encoded_cells: list[str] = []
        for span in row:
            attribute = "" if span is None else f' w:val="{span}"'
            encoded_cells.append(f"<w:tc><w:tcPr><w:gridSpan{attribute}/></w:tcPr><w:p/></w:tc>")
        encoded_rows.append(f"<w:tr>{''.join(encoded_cells)}</w:tr>")
    return (
        f'<w:document xmlns:w="{WORD}"><w:body><w:tbl>'
        f"{''.join(encoded_rows)}</w:tbl></w:body></w:document>"
    ).encode()


def _vertical_geometry_xml(rows: int) -> bytes:
    encoded_rows = []
    for index in range(rows):
        merge = ' w:val="restart"' if index == 0 else ""
        encoded_rows.append(f"<w:tr><w:tc><w:tcPr><w:vMerge{merge}/></w:tcPr><w:p/></w:tc></w:tr>")
    return (
        f'<w:document xmlns:w="{WORD}"><w:body><w:tbl>'
        f"{''.join(encoded_rows)}</w:tbl></w:body></w:document>"
    ).encode()


def _docx_with_grid_span(span: int) -> bytes:
    source = BytesIO()
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Synthetic"
    document.save(source)

    with ZipFile(BytesIO(source.getvalue())) as archive:
        members = {member.filename: archive.read(member) for member in archive.infolist()}
    document_xml = members["word/document.xml"]
    members["word/document.xml"] = document_xml.replace(
        b"<w:tcPr>",
        f'<w:tcPr><w:gridSpan w:val="{span}"/>'.encode(),
        1,
    )
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        for name, value in members.items():
            archive.writestr(name, value)
    return output.getvalue()


def _docx_with_vertical_merge(rows: int) -> bytes:
    stream = BytesIO()
    document = Document()
    table = document.add_table(rows=rows, cols=1)
    merged = table.cell(0, 0).merge(table.cell(rows - 1, 0))
    merged.text = "Synthetic vertical assessment"
    document.save(stream)
    return stream.getvalue()
