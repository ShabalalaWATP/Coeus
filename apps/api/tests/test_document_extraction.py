from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from reportlab.pdfgen.canvas import Canvas

from coeus.services import document_extraction
from coeus.services.document_extraction import DocumentExtractionError, extract_pages


def test_extracts_pdf_pages_for_citation() -> None:
    stream = BytesIO()
    canvas = Canvas(stream)
    canvas.drawString(72, 720, "Synthetic maritime movement assessment")
    canvas.showPage()
    canvas.drawString(72, 720, "Second page evidence")
    canvas.save()

    pages = extract_pages(stream.getvalue(), "application/pdf")

    assert [page.page_number for page in pages] == [1, 2]
    assert "maritime movement" in pages[0].text
    assert "Second page" in pages[1].text


def test_extracts_docx_paragraphs_and_tables() -> None:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("Synthetic intelligence report")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Location"
    table.cell(0, 1).text = "Mock harbour"
    document.save(stream)

    pages = extract_pages(
        stream.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert len(pages) == 1
    assert "Synthetic intelligence report" in pages[0].text
    assert "Mock harbour" in pages[0].text


@pytest.mark.parametrize(
    ("content", "mime_type", "code"),
    [
        (b"not a pdf", "application/pdf", "pdf_signature_invalid"),
        (
            b"not a docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx_signature_invalid",
        ),
        (b"plain", "text/plain", "asset_type_unsupported"),
        (b"", "application/pdf", "asset_size_invalid"),
    ],
)
def test_rejects_malformed_or_unsupported_assets(content: bytes, mime_type: str, code: str) -> None:
    with pytest.raises(DocumentExtractionError, match=code):
        extract_pages(content, mime_type)


def test_rejects_macro_enabled_docx_content() -> None:
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", "<document />")
        archive.writestr("word/vbaProject.bin", b"macro")

    with pytest.raises(DocumentExtractionError, match="docx_not_extractable"):
        extract_pages(
            stream.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


def test_rejects_pdf_page_limit_and_malformed_pdf(monkeypatch: pytest.MonkeyPatch) -> None:
    stream = BytesIO()
    canvas = Canvas(stream)
    canvas.drawString(72, 720, "Synthetic page")
    canvas.save()

    monkeypatch.setattr(document_extraction, "MAX_DOCUMENT_PAGES", 0)
    with pytest.raises(DocumentExtractionError, match="pdf_not_extractable"):
        extract_pages(stream.getvalue(), "application/pdf")
    with pytest.raises(DocumentExtractionError, match="pdf_parse_failed"):
        extract_pages(b"%PDF-invalid", "application/pdf")


def test_rejects_docx_archive_limits_and_missing_document(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", "<document />")
    content = stream.getvalue()

    monkeypatch.setattr(document_extraction, "MAX_ZIP_MEMBERS", 0)
    with pytest.raises(DocumentExtractionError, match="docx_archive_too_large"):
        extract_pages(
            content,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    monkeypatch.setattr(document_extraction, "MAX_ZIP_MEMBERS", 5_000)
    monkeypatch.setattr(document_extraction, "MAX_ZIP_EXPANDED_BYTES", 0)
    with pytest.raises(DocumentExtractionError, match="docx_archive_too_large"):
        extract_pages(
            content,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    missing = BytesIO()
    with ZipFile(missing, "w", ZIP_DEFLATED) as archive:
        archive.writestr("word/styles.xml", "<styles />")
    monkeypatch.setattr(document_extraction, "MAX_ZIP_EXPANDED_BYTES", 50_000_000)
    with pytest.raises(DocumentExtractionError, match="docx_not_extractable"):
        extract_pages(
            missing.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    with pytest.raises(DocumentExtractionError, match="docx_parse_failed"):
        extract_pages(
            b"PK-invalid",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


def test_bounds_extracted_text_and_discards_empty_pages(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pages = (
        document_extraction.ExtractedPage(1, ""),
        document_extraction.ExtractedPage(2, "long text"),
    )
    monkeypatch.setattr(document_extraction, "MAX_EXTRACTED_CHARACTERS", 1)
    with pytest.raises(DocumentExtractionError, match="extracted_text_too_large"):
        document_extraction._bounded_pages(pages)

    assert document_extraction._normalise("one\x00  two\n\n\nthree") == "one two\n\nthree"


def test_pptx_xml_rejects_entities_before_parser_expansion() -> None:
    xml = b"""<?xml version="1.0"?>
    <!DOCTYPE p [<!ENTITY amplified "synthetic synthetic synthetic">]>
    <p:sld xmlns:p="urn:p" xmlns:a="urn:a"><a:t>&amplified;</a:t></p:sld>"""

    with pytest.raises(DocumentExtractionError, match="pptx_xml_unsafe"):
        document_extraction._pptx_text(xml)


def test_pptx_extraction_rejects_invalid_structure_and_bounds_pages(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    with pytest.raises(DocumentExtractionError, match="pptx_signature_invalid"):
        extract_pages(b"not a presentation", document_extraction.PPTX_MIME)

    missing_presentation = _pptx_archive({"ppt/slides/slide1.xml": "<p:sld xmlns:p='urn:p' />"})
    with pytest.raises(DocumentExtractionError, match="pptx_not_extractable"):
        extract_pages(missing_presentation, document_extraction.PPTX_MIME)

    missing_slides = _pptx_archive({"ppt/presentation.xml": "<p:presentation xmlns:p='urn:p' />"})
    with pytest.raises(DocumentExtractionError, match="pptx_not_extractable"):
        extract_pages(missing_slides, document_extraction.PPTX_MIME)

    one_slide = _pptx_archive(
        {
            "ppt/presentation.xml": "<p:presentation xmlns:p='urn:p' />",
            "ppt/slides/slide1.xml": (
                "<p:sld xmlns:p='urn:p' xmlns:a='urn:a'><a:t>Synthetic slide</a:t></p:sld>"
            ),
        }
    )
    monkeypatch.setattr(document_extraction, "MAX_DOCUMENT_PAGES", 0)
    with pytest.raises(DocumentExtractionError, match="pptx_not_extractable"):
        extract_pages(one_slide, document_extraction.PPTX_MIME)

    monkeypatch.setattr(document_extraction, "MAX_DOCUMENT_PAGES", 200)
    pages = extract_pages(one_slide, document_extraction.PPTX_MIME)
    assert [(page.page_number, page.text) for page in pages] == [(1, "Synthetic slide")]


def _pptx_archive(parts: dict[str, str]) -> bytes:
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        for name, content in parts.items():
            archive.writestr(name, content)
    return stream.getvalue()
