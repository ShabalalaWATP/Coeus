"""Bounded, non-executing text extraction for supported search assets."""

from dataclasses import dataclass
from io import BytesIO
from re import sub
from zipfile import ZipFile, ZipInfo

from defusedxml import ElementTree
from defusedxml.common import DefusedXmlException
from pypdf import PdfReader

from coeus.services.docx_geometry import (
    DocxTableGeometryError,
    validate_docx_table_geometry,
)
from coeus.services.office_archive import (
    OfficeArchiveInvalidError,
    OfficeArchiveLimitError,
    preflight_office_archive,
    read_bounded_member,
    validate_office_archive,
)
from coeus.services.pdf_decode_budget import (
    PdfDecodedStreamLimitError,
    validate_pdf_decoded_streams,
)

MAX_DOCUMENT_BYTES = 20_000_000
MAX_DOCUMENT_PAGES = 200
MAX_EXTRACTED_CHARACTERS = 2_000_000
MAX_ZIP_MEMBERS = 5_000
MAX_ZIP_EXPANDED_BYTES = 50_000_000
MAX_ZIP_CENTRAL_DIRECTORY_BYTES = MAX_DOCUMENT_BYTES
MAX_XML_PART_BYTES = 5_000_000
EXTRACTOR_VERSION = "local-pdf-docx-pptx-v2"
PPTX_MIME = "application/vnd.openxmlformats-officedocument.presentationml.presentation"


class DocumentExtractionError(ValueError):
    def __init__(self, code: str) -> None:
        super().__init__(code)
        self.code = code


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int
    text: str


def extract_pages(content: bytes, mime_type: str) -> tuple[ExtractedPage, ...]:
    if not content or len(content) > MAX_DOCUMENT_BYTES:
        raise DocumentExtractionError("asset_size_invalid")
    if mime_type == "application/pdf":
        return _extract_pdf(content)
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(content)
    if mime_type == PPTX_MIME:
        return _extract_pptx(content)
    raise DocumentExtractionError("asset_type_unsupported")


def _extract_pdf(content: bytes) -> tuple[ExtractedPage, ...]:
    if not content.startswith(b"%PDF-"):
        raise DocumentExtractionError("pdf_signature_invalid")
    try:
        reader = PdfReader(BytesIO(content), strict=True)
        if reader.is_encrypted or len(reader.pages) > MAX_DOCUMENT_PAGES:
            raise DocumentExtractionError("pdf_not_extractable")
        validate_pdf_decoded_streams(reader.pages)
        total = 0
        retained: list[ExtractedPage] = []
        for index, page in enumerate(reader.pages, start=1):
            text = _normalise(page.extract_text() or "")
            if not text:
                continue
            total += len(text)
            if total > MAX_EXTRACTED_CHARACTERS:
                raise DocumentExtractionError("extracted_text_too_large")
            retained.append(ExtractedPage(index, text))
    except DocumentExtractionError:
        raise
    except PdfDecodedStreamLimitError as exc:
        raise DocumentExtractionError("pdf_processing_limit_exceeded") from exc
    except Exception as exc:
        raise DocumentExtractionError("pdf_parse_failed") from exc
    return tuple(retained)


def _extract_docx(content: bytes) -> tuple[ExtractedPage, ...]:
    if not content.startswith(b"PK"):
        raise DocumentExtractionError("docx_signature_invalid")
    try:
        _preflight_office_archive(content, "docx")
        with ZipFile(BytesIO(content)) as archive:
            members = _validated_office_members(archive, "docx")
            names = {member.filename.casefold() for member in members}
            if "word/vbaproject.bin" in names or "word/document.xml" not in names:
                raise DocumentExtractionError("docx_not_extractable")
            document_members = tuple(
                member for member in members if member.filename.casefold() == "word/document.xml"
            )
            if len(document_members) != 1:
                raise DocumentExtractionError("docx_not_extractable")
            document_member = document_members[0]
            document_xml = read_bounded_member(archive, document_member, MAX_XML_PART_BYTES)
            _validate_docx_table_geometry(document_xml)
        from docx import Document

        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        ]
        text = _normalise("\n".join((*paragraphs, *table_cells)))
    except DocumentExtractionError:
        raise
    except OfficeArchiveLimitError as exc:
        raise DocumentExtractionError("docx_archive_too_large") from exc
    except Exception as exc:
        raise DocumentExtractionError("docx_parse_failed") from exc
    return _bounded_pages((ExtractedPage(1, text),))


def _extract_pptx(content: bytes) -> tuple[ExtractedPage, ...]:
    if not content.startswith(b"PK"):
        raise DocumentExtractionError("pptx_signature_invalid")
    try:
        _preflight_office_archive(content, "pptx")
        with ZipFile(BytesIO(content)) as archive:
            members = _validated_office_members(archive, "pptx")
            names = {member.filename.casefold() for member in members}
            if "ppt/vbaproject.bin" in names or "ppt/presentation.xml" not in names:
                raise DocumentExtractionError("pptx_not_extractable")
            slide_members = sorted(
                (
                    member
                    for member in members
                    if member.filename.startswith("ppt/slides/slide")
                    and member.filename.endswith(".xml")
                ),
                key=lambda member: _slide_number(member.filename),
            )
            if not slide_members or len(slide_members) > MAX_DOCUMENT_PAGES:
                raise DocumentExtractionError("pptx_not_extractable")
            pages = tuple(
                ExtractedPage(
                    index,
                    _pptx_text(read_bounded_member(archive, member, MAX_XML_PART_BYTES)),
                )
                for index, member in enumerate(slide_members, start=1)
            )
    except DocumentExtractionError:
        raise
    except Exception as exc:
        raise DocumentExtractionError("pptx_parse_failed") from exc
    return _bounded_pages(pages)


def _validated_office_members(archive: ZipFile, prefix: str) -> tuple[ZipInfo, ...]:
    try:
        return validate_office_archive(
            archive,
            max_members=MAX_ZIP_MEMBERS,
            max_expanded_bytes=MAX_ZIP_EXPANDED_BYTES,
        )
    except OfficeArchiveLimitError as exc:
        raise DocumentExtractionError(f"{prefix}_archive_too_large") from exc


def _preflight_office_archive(content: bytes, prefix: str) -> None:
    try:
        preflight_office_archive(
            content,
            max_members=MAX_ZIP_MEMBERS,
            max_directory_bytes=MAX_ZIP_CENTRAL_DIRECTORY_BYTES,
        )
    except OfficeArchiveLimitError as exc:
        raise DocumentExtractionError(f"{prefix}_archive_too_large") from exc
    except OfficeArchiveInvalidError as exc:
        raise DocumentExtractionError(f"{prefix}_parse_failed") from exc


def _slide_number(name: str) -> int:
    stem = name.rsplit("/", 1)[-1].removeprefix("slide").removesuffix(".xml")
    return int(stem)


def _pptx_text(xml: bytes) -> str:
    try:
        root = ElementTree.fromstring(xml, forbid_dtd=True)
    except DefusedXmlException as exc:
        raise DocumentExtractionError("pptx_xml_unsafe") from exc
    return _normalise("\n".join(node.text or "" for node in root.iter() if node.tag.endswith("}t")))


def _validate_docx_table_geometry(xml: bytes) -> None:
    try:
        validate_docx_table_geometry(xml)
    except DocxTableGeometryError as exc:
        raise DocumentExtractionError(exc.code) from exc


def _bounded_pages(pages: tuple[ExtractedPage, ...]) -> tuple[ExtractedPage, ...]:
    total = 0
    retained: list[ExtractedPage] = []
    for page in pages:
        if not page.text:
            continue
        total += len(page.text)
        if total > MAX_EXTRACTED_CHARACTERS:
            raise DocumentExtractionError("extracted_text_too_large")
        retained.append(page)
    return tuple(retained)


def _normalise(text: str) -> str:
    without_controls = "".join(
        character for character in text if character in "\n\t" or ord(character) >= 32
    )
    return sub(r"[ \t]+", " ", sub(r"\n{3,}", "\n\n", without_controls)).strip()
